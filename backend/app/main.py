from __future__ import annotations

import asyncio
import hashlib
import json
import os
import re
import shutil
import uuid
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Optional

from docx import Document as DocxDocument
from fastapi import Depends, FastAPI, File, Form, Header, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from openai import OpenAI
from pydantic import BaseModel
from pypdf import PdfReader
from sqlalchemy import (
    Boolean,
    Date,
    DateTime,
    ForeignKey,
    Integer,
    String,
    Text,
    create_engine,
    func,
)
from sqlalchemy.orm import DeclarativeBase, Mapped, Session, mapped_column, relationship, sessionmaker


DATABASE_URL = os.getenv(
    "DATABASE_URL",
    "postgresql+psycopg2://senior_qa:local_dev_value@localhost:5432/senior_qa",
)
UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "/data/uploads"))
ADMIN_USERNAME = os.getenv("ADMIN_USERNAME", "admin")
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "replace_me_before_use")
ADMIN_FULL_NAME = os.getenv("ADMIN_FULL_NAME", "Senior QA Admin")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4.1-mini")
MAX_RETRIEVAL_CHUNKS = int(os.getenv("MAX_RETRIEVAL_CHUNKS", "8"))
AUTONOMOUS_WORKER_ENABLED = os.getenv("AUTONOMOUS_WORKER_ENABLED", "true").lower() == "true"
AUTONOMOUS_WORKER_INTERVAL_SECONDS = int(os.getenv("AUTONOMOUS_WORKER_INTERVAL_SECONDS", "300"))


class Base(DeclarativeBase):
    pass


class User(Base):
    __tablename__ = "users"

    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    username: Mapped[str] = mapped_column(String(120), unique=True, index=True)
    full_name: Mapped[str] = mapped_column(String(255), default="")
    password_hash: Mapped[str] = mapped_column(String(255))
    role: Mapped[str] = mapped_column(String(50), default="QA")
    is_active: Mapped[bool] = mapped_column(Boolean, default=True)
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=lambda: utcnow())


class UserSession(Base):
    __tablename__ = "user_sessions"

    token: Mapped[str] = mapped_column(String(255), primary_key=True)
    user_id: Mapped[int] = mapped_column(ForeignKey("users.id"), index=True)
    expires_at: Mapped[datetime] = mapped_column(DateTime(timezone=True))
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=lambda: utcnow())
    user: Mapped[User] = relationship()


class KnowledgeDocument(Base):
    __tablename__ = "knowledge_documents"

    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    title: Mapped[str] = mapped_column(String(500), index=True)
    document_number: Mapped[str] = mapped_column(String(120), default="", index=True)
    version: Mapped[str] = mapped_column(String(50), default="")
    status: Mapped[str] = mapped_column(String(40), default="current", index=True)
    doc_type: Mapped[str] = mapped_column(String(80), default="SOP", index=True)
    department: Mapped[str] = mapped_column(String(120), default="")
    effective_date: Mapped[Optional[date]] = mapped_column(Date, nullable=True)
    review_date: Mapped[Optional[date]] = mapped_column(Date, nullable=True)
    source_filename: Mapped[str] = mapped_column(String(500), default="")
    file_path: Mapped[str] = mapped_column(String(1000), default="")
    uploaded_by_id: Mapped[Optional[int]] = mapped_column(ForeignKey("users.id"), nullable=True)
    uploaded_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=lambda: utcnow())
    chunks: Mapped[list["DocumentChunk"]] = relationship(
        back_populates="document", cascade="all, delete-orphan"
    )


class DocumentChunk(Base):
    __tablename__ = "document_chunks"

    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    document_id: Mapped[int] = mapped_column(ForeignKey("knowledge_documents.id"), index=True)
    chunk_index: Mapped[int] = mapped_column(Integer, default=0)
    content: Mapped[str] = mapped_column(Text)
    document: Mapped[KnowledgeDocument] = relationship(back_populates="chunks")


class QmsTask(Base):
    __tablename__ = "qms_tasks"

    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    task_code: Mapped[str] = mapped_column(String(50), unique=True, index=True)
    title: Mapped[str] = mapped_column(String(500))
    task_type: Mapped[str] = mapped_column(String(80), index=True)
    description: Mapped[str] = mapped_column(Text, default="")
    priority: Mapped[str] = mapped_column(String(30), default="Medium")
    status: Mapped[str] = mapped_column(String(50), default="New", index=True)
    due_date: Mapped[Optional[date]] = mapped_column(Date, nullable=True)
    created_by_id: Mapped[Optional[int]] = mapped_column(ForeignKey("users.id"), nullable=True)
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=lambda: utcnow())
    updated_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=lambda: utcnow())
    last_agent_run_at: Mapped[Optional[datetime]] = mapped_column(DateTime(timezone=True), nullable=True)
    questions: Mapped[list["AgentQuestion"]] = relationship(
        back_populates="task", cascade="all, delete-orphan"
    )
    outputs: Mapped[list["TaskOutput"]] = relationship(
        back_populates="task", cascade="all, delete-orphan"
    )


class AgentQuestion(Base):
    __tablename__ = "agent_questions"

    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    task_id: Mapped[int] = mapped_column(ForeignKey("qms_tasks.id"), index=True)
    question: Mapped[str] = mapped_column(Text)
    assigned_to: Mapped[str] = mapped_column(String(255), default="QA")
    status: Mapped[str] = mapped_column(String(50), default="Open", index=True)
    answer: Mapped[str] = mapped_column(Text, default="")
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=lambda: utcnow())
    answered_at: Mapped[Optional[datetime]] = mapped_column(DateTime(timezone=True), nullable=True)
    task: Mapped[QmsTask] = relationship(back_populates="questions")


class TaskOutput(Base):
    __tablename__ = "task_outputs"

    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    task_id: Mapped[int] = mapped_column(ForeignKey("qms_tasks.id"), index=True)
    title: Mapped[str] = mapped_column(String(500))
    content: Mapped[str] = mapped_column(Text)
    sources_json: Mapped[str] = mapped_column(Text, default="[]")
    assumptions_json: Mapped[str] = mapped_column(Text, default="[]")
    status: Mapped[str] = mapped_column(String(50), default="Draft Prepared", index=True)
    reviewer_comment: Mapped[str] = mapped_column(Text, default="")
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=lambda: utcnow())
    reviewed_at: Mapped[Optional[datetime]] = mapped_column(DateTime(timezone=True), nullable=True)
    task: Mapped[QmsTask] = relationship(back_populates="outputs")


class AgentRun(Base):
    __tablename__ = "agent_runs"

    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    task_id: Mapped[Optional[int]] = mapped_column(ForeignKey("qms_tasks.id"), nullable=True, index=True)
    run_type: Mapped[str] = mapped_column(String(80), default="manual")
    status: Mapped[str] = mapped_column(String(50), default="started")
    summary: Mapped[str] = mapped_column(Text, default="")
    sources_json: Mapped[str] = mapped_column(Text, default="[]")
    started_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=lambda: utcnow())
    finished_at: Mapped[Optional[datetime]] = mapped_column(DateTime(timezone=True), nullable=True)


class AuditEvent(Base):
    __tablename__ = "audit_events"

    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    event_type: Mapped[str] = mapped_column(String(120), index=True)
    entity_type: Mapped[str] = mapped_column(String(120), default="")
    entity_id: Mapped[str] = mapped_column(String(120), default="")
    user_id: Mapped[Optional[int]] = mapped_column(ForeignKey("users.id"), nullable=True)
    details_json: Mapped[str] = mapped_column(Text, default="{}")
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=lambda: utcnow())


engine = create_engine(DATABASE_URL, pool_pre_ping=True)
SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)


def utcnow() -> datetime:
    return datetime.now(timezone.utc)


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


class LoginIn(BaseModel):
    username: str
    password: str


class TokenOut(BaseModel):
    access_token: str
    token_type: str = "bearer"


class TaskCreate(BaseModel):
    title: str
    task_type: str
    description: str = ""
    priority: str = "Medium"
    due_date: Optional[date] = None


class TaskUpdate(BaseModel):
    title: Optional[str] = None
    task_type: Optional[str] = None
    description: Optional[str] = None
    priority: Optional[str] = None
    due_date: Optional[date] = None
    status: Optional[str] = None


class AnswerQuestionIn(BaseModel):
    answer: str


class ReviewOutputIn(BaseModel):
    decision: str
    comment: str = ""


class SearchRequest(BaseModel):
    query: str
    include_superseded: bool = False


def hash_password(password: str) -> str:
    salt = uuid.uuid4().hex
    digest = hashlib.sha256((salt + password).encode("utf-8")).hexdigest()
    return f"{salt}:{digest}"


def verify_password(password: str, stored_hash: str) -> bool:
    try:
        salt, digest = stored_hash.split(":", 1)
    except ValueError:
        return False
    expected = hashlib.sha256((salt + password).encode("utf-8")).hexdigest()
    return expected == digest


def create_session(db: Session, user: User) -> str:
    token = uuid.uuid4().hex + uuid.uuid4().hex
    session = UserSession(
        token=token,
        user_id=user.id,
        expires_at=utcnow() + timedelta(minutes=int(os.getenv("ACCESS_TOKEN_EXPIRE_MINUTES", "720"))),
    )
    db.add(session)
    db.commit()
    return token


def get_current_user(
    authorization: Optional[str] = Header(default=None),
    db: Session = Depends(get_db),
) -> User:
    if not authorization or not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Missing bearer token")
    token = authorization.split(" ", 1)[1].strip()
    session = db.get(UserSession, token)
    if not session or session.expires_at < utcnow():
        raise HTTPException(status_code=401, detail="Invalid or expired session")
    user = db.get(User, session.user_id)
    if not user or not user.is_active:
        raise HTTPException(status_code=401, detail="User inactive")
    return user


def audit(
    db: Session,
    event_type: str,
    entity_type: str = "",
    entity_id: str | int = "",
    user: Optional[User] = None,
    details: Optional[dict[str, Any]] = None,
) -> None:
    db.add(
        AuditEvent(
            event_type=event_type,
            entity_type=entity_type,
            entity_id=str(entity_id or ""),
            user_id=user.id if user else None,
            details_json=json.dumps(details or {}, default=str),
        )
    )
    db.commit()


def parse_date_or_none(value: str | None) -> Optional[date]:
    if not value:
        return None
    return date.fromisoformat(value)


def extract_text_from_upload(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(str(file_path))
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        return "\n\n".join(pages)
    if suffix == ".docx":
        doc = DocxDocument(str(file_path))
        return "\n".join(p.text for p in doc.paragraphs)
    if suffix in {".txt", ".md", ".csv"}:
        return file_path.read_text(encoding="utf-8", errors="ignore")
    return file_path.read_text(encoding="utf-8", errors="ignore")


def chunk_text(text: str, max_chars: int = 1800, overlap: int = 250) -> list[str]:
    cleaned = re.sub(r"\s+", " ", text).strip()
    if not cleaned:
        return []
    chunks = []
    start = 0
    while start < len(cleaned):
        end = min(start + max_chars, len(cleaned))
        chunks.append(cleaned[start:end])
        if end == len(cleaned):
            break
        start = max(0, end - overlap)
    return chunks


STOPWORDS = {
    "the", "and", "for", "with", "that", "this", "from", "into", "your", "have", "will",
    "shall", "should", "must", "are", "was", "were", "been", "being", "task", "draft",
}


def tokens(text: str) -> set[str]:
    return {t for t in re.findall(r"[a-zA-Z0-9]{3,}", text.lower()) if t not in STOPWORDS}


def search_chunks(db: Session, query: str, include_superseded: bool = False, limit: int = MAX_RETRIEVAL_CHUNKS):
    q_tokens = tokens(query)
    rows = db.query(DocumentChunk).join(KnowledgeDocument).all()
    scored = []
    for chunk in rows:
        if not include_superseded and chunk.document.status.lower() != "current":
            continue
        c_tokens = tokens(chunk.content + " " + chunk.document.title + " " + chunk.document.document_number)
        score = len(q_tokens & c_tokens)
        if score > 0:
            scored.append((score, chunk))
    scored.sort(key=lambda item: item[0], reverse=True)
    return [chunk for _, chunk in scored[:limit]]


def source_payload(chunks: list[DocumentChunk]) -> list[dict[str, Any]]:
    return [
        {
            "document_id": c.document.id,
            "title": c.document.title,
            "document_number": c.document.document_number,
            "version": c.document.version,
            "status": c.document.status,
            "chunk_index": c.chunk_index,
        }
        for c in chunks
    ]


def task_context(task: QmsTask) -> str:
    return f"{task.title}\n{task.task_type}\n{task.description}"


def open_questions(db: Session, task_id: int) -> list[AgentQuestion]:
    return (
        db.query(AgentQuestion)
        .filter(AgentQuestion.task_id == task_id, AgentQuestion.status == "Open")
        .order_by(AgentQuestion.created_at.asc())
        .all()
    )


def answered_questions(db: Session, task_id: int) -> list[AgentQuestion]:
    return (
        db.query(AgentQuestion)
        .filter(AgentQuestion.task_id == task_id, AgentQuestion.status == "Answered")
        .order_by(AgentQuestion.created_at.asc())
        .all()
    )


def add_unique_question(db: Session, task: QmsTask, question: str, assigned_to: str = "QA") -> None:
    existing = (
        db.query(AgentQuestion)
        .filter(AgentQuestion.task_id == task.id, AgentQuestion.question == question)
        .first()
    )
    if existing:
        return
    db.add(AgentQuestion(task_id=task.id, question=question, assigned_to=assigned_to))
    task.status = "Information Required"
    task.updated_at = utcnow()
    db.commit()


def missing_information_questions(task: QmsTask) -> list[str]:
    text = task_context(task).lower()
    questions: list[str] = []

    if len(text.strip()) < 120:
        questions.append("Please add more background detail for this task, including what happened, why the task is required, and what output you expect.")

    if "deviation" in task.task_type.lower() or "deviation" in text:
        if "batch" not in text:
            questions.append("Please confirm the affected product code, batch number, and whether any related batches may be impacted.")
        if "impact" not in text and "quality" not in text:
            questions.append("Please provide the initial product quality impact assessment or confirm that this still needs to be drafted.")
        if "root cause" not in text and "cause" not in text:
            questions.append("Please confirm any known or suspected root cause, or state that root cause is currently unknown.")

    if "supplier" in task.task_type.lower() or "supplier" in text:
        if "coa" not in text and "certificate" not in text:
            questions.append("Please confirm what testing evidence is available, including CoA, ID testing, full testing, and manufacturer qualification evidence.")
        if "manufacturer" not in text:
            questions.append("Please confirm the original manufacturer and whether the distributor repacks, relabels, imports, or only distributes the material.")

    if "sop" in task.task_type.lower() or "procedure" in task.task_type.lower():
        if "compare" not in text and "review" not in text and "gap" not in text:
            questions.append("Please confirm whether the task is a gap assessment, rewrite, periodic review, or new SOP draft.")

    return questions[:5]


def fallback_draft(task: QmsTask, chunks: list[DocumentChunk], answers: list[AgentQuestion]) -> str:
    source_lines = "\n".join(
        f"- {c.document.document_number or 'N/A'} {c.document.title} v{c.document.version}, chunk {c.chunk_index}"
        for c in chunks
    ) or "- No directly relevant knowledge-base documents were retrieved."

    answer_lines = "\n".join(f"- Q: {q.question}\n  A: {q.answer}" for q in answers) or "- No answered clarification questions were available."

    return f"""# Draft QA Output — {task.title}

**Task type:** {task.task_type}  
**Task code:** {task.task_code}  
**Status:** Draft prepared for human review only.

## 1. Draft summary

This is a controlled draft prepared by the Senior QA AI assistant based on the task description, answered clarification questions, and the uploaded knowledge-base documents retrieved for this task.

## 2. Information reviewed

{source_lines}

## 3. Clarification answers considered

{answer_lines}

## 4. Draft QA assessment

Based on the available information, the record should be reviewed against the relevant internal SOPs/forms and any applicable GMP/GDP/MHRA/ICH expectations contained in the uploaded knowledge base.

The human reviewer should confirm:

- Whether the correct current internal procedure/form has been used.
- Whether the facts are complete, accurate, attributable, legible, contemporaneous, original/true-copy where applicable, and accurate.
- Whether product quality, patient safety, regulatory compliance, and data integrity impact have been assessed.
- Whether any CAPA, change control, training, supplier action, or document update is required.
- Whether the proposed wording is suitable for the live QMS record.

## 5. Draft output / proposed wording

Use the following as a starting point only:

> A review has been performed against the available information and relevant controlled copies of internal QMS documentation. The event/task has been assessed for potential impact to product quality, patient safety, regulatory compliance and data integrity. The available information suggests that further human QA review is required before any final conclusion, closure or approval decision is made.

## 6. Missing information / reviewer checks

- Confirm all facts against the live approved QMS record.
- Confirm document versions are current copies.
- Confirm whether additional evidence is required.
- Confirm whether the draft is suitable for approval, rework, or rejection.

## 7. Human review requirement

This output is not approved. A competent QA reviewer must review, amend as needed, and approve or reject the output in accordance with the applicable QMS process.
"""


async def ai_draft(task: QmsTask, chunks: list[DocumentChunk], answers: list[AgentQuestion]) -> str:
    if not OPENAI_API_KEY:
        return fallback_draft(task, chunks, answers)

    client = OpenAI(api_key=OPENAI_API_KEY)
    sources = "\n\n".join(
        f"SOURCE {idx + 1}: {c.document.document_number} {c.document.title} v{c.document.version}, chunk {c.chunk_index}\n{c.content}"
        for idx, c in enumerate(chunks)
    )
    clarifications = "\n".join(f"Q: {q.question}\nA: {q.answer}" for q in answers)

    prompt = f"""
You are an experienced UK pharmaceutical Quality Assurance Officer.
You are preparing a draft QMS-support output for human review only.

Hard rules:
- Do not approve, release, reject, close, or make final GMP decisions.
- Use only the task details, answered questions, and provided source excerpts.
- Make assumptions explicit.
- Include missing information.
- Write in formal QA/GMP style.
- Include a source section listing the documents relied upon.

Task:
{task.task_code} — {task.title}
Type: {task.task_type}
Priority: {task.priority}
Description:
{task.description}

Answered clarification questions:
{clarifications or "None"}

Retrieved source excerpts:
{sources or "No source excerpts retrieved."}

Prepare:
1. Draft summary
2. Applicable source basis
3. Draft QA assessment
4. Product quality / patient safety / compliance / data integrity considerations where relevant
5. Proposed QMS wording
6. Missing information and assumptions
7. Human review checklist
"""

    response = client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=[
            {"role": "system", "content": "You are a controlled AI QA drafting assistant for pharmaceutical QMS work."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
    )
    return response.choices[0].message.content or ""


async def run_agent_for_task(db: Session, task_id: int, run_type: str = "manual") -> AgentRun:
    task = db.get(QmsTask, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")

    run = AgentRun(task_id=task.id, run_type=run_type, status="started")
    db.add(run)
    db.commit()
    db.refresh(run)

    try:
        if open_questions(db, task.id):
            task.status = "Information Required"
            task.updated_at = utcnow()
            run.status = "blocked"
            run.summary = "Task has open clarification questions."
            run.finished_at = utcnow()
            db.commit()
            return run

        chunks = search_chunks(db, task_context(task), limit=MAX_RETRIEVAL_CHUNKS)
        if not chunks:
            add_unique_question(
                db,
                task,
                "No current knowledge-base documents were retrieved for this task. Please upload the relevant SOP/form/guidance or add more task detail so the AI can retrieve the correct documents.",
                "QA/Admin",
            )
            run.status = "blocked"
            run.summary = "No relevant current knowledge-base documents were found."
            run.finished_at = utcnow()
            db.commit()
            return run

        for question in missing_information_questions(task):
            add_unique_question(db, task, question, "Task Owner")

        if open_questions(db, task.id):
            run.status = "blocked"
            run.summary = "The AI generated missing-information questions before drafting."
            run.sources_json = json.dumps(source_payload(chunks), default=str)
            run.finished_at = utcnow()
            db.commit()
            return run

        answers = answered_questions(db, task.id)
        draft = await ai_draft(task, chunks, answers)

        output = TaskOutput(
            task_id=task.id,
            title=f"AI draft output for {task.task_code}",
            content=draft,
            sources_json=json.dumps(source_payload(chunks), default=str),
            assumptions_json=json.dumps(
                [
                    "The output is draft-only and requires competent human QA review.",
                    "Uploaded copies are assumed to represent the current controlled documents where status is set to current.",
                    "The live QMS remains the system of record.",
                ],
                default=str,
            ),
            status="Draft Prepared",
        )
        db.add(output)

        task.status = "Awaiting Human Review"
        task.updated_at = utcnow()
        task.last_agent_run_at = utcnow()
        run.status = "completed"
        run.summary = "Draft output prepared and routed for human review."
        run.sources_json = json.dumps(source_payload(chunks), default=str)
        run.finished_at = utcnow()
        db.commit()
        return run
    except Exception as exc:
        task.status = "Agent Error"
        task.updated_at = utcnow()
        run.status = "error"
        run.summary = str(exc)
        run.finished_at = utcnow()
        db.commit()
        raise


async def autonomous_worker_loop() -> None:
    while True:
        await asyncio.sleep(AUTONOMOUS_WORKER_INTERVAL_SECONDS)
        db = SessionLocal()
        try:
            candidates = (
                db.query(QmsTask)
                .filter(QmsTask.status.in_(["New", "Assigned to AI", "In Progress", "Returned for Rework"]))
                .order_by(QmsTask.due_date.asc().nullslast(), QmsTask.created_at.asc())
                .limit(10)
                .all()
            )
            for task in candidates:
                try:
                    task.status = "In Progress"
                    task.updated_at = utcnow()
                    db.commit()
                    await run_agent_for_task(db, task.id, run_type="autonomous")
                except Exception:
                    db.rollback()
        finally:
            db.close()


app = FastAPI(title="Senior QA AI Officer", version="0.3.0-phase3")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
async def startup() -> None:
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    Base.metadata.create_all(bind=engine)

    db = SessionLocal()
    try:
        admin = db.query(User).filter(User.username == ADMIN_USERNAME).first()
        if not admin:
            db.add(
                User(
                    username=ADMIN_USERNAME,
                    full_name=ADMIN_FULL_NAME,
                    password_hash=hash_password(ADMIN_PASSWORD),
                    role="ADMIN",
                    is_active=True,
                )
            )
            db.commit()
        if AUTONOMOUS_WORKER_ENABLED:
            asyncio.create_task(autonomous_worker_loop())
    finally:
        db.close()


@app.get("/health")
def health() -> dict[str, Any]:
    return {
        "status": "ok",
        "app": "Senior QA AI Officer",
        "phase": "3",
        "autonomous_worker_enabled": AUTONOMOUS_WORKER_ENABLED,
        "openai_configured": bool(OPENAI_API_KEY),
    }


@app.post("/auth/login", response_model=TokenOut)
def login(payload: LoginIn, db: Session = Depends(get_db)) -> TokenOut:
    user = db.query(User).filter(User.username == payload.username).first()
    if not user or not verify_password(payload.password, user.password_hash):
        audit(db, "LOGIN_FAILED", "user", payload.username, None, {"username": payload.username})
        raise HTTPException(status_code=401, detail="Invalid username or password")
    if not user.is_active:
        raise HTTPException(status_code=401, detail="User inactive")
    token = create_session(db, user)
    audit(db, "LOGIN_SUCCESS", "user", user.id, user, {"username": user.username})
    return TokenOut(access_token=token)


@app.get("/auth/me")
def me(current_user: User = Depends(get_current_user)) -> dict[str, Any]:
    return {
        "id": current_user.id,
        "username": current_user.username,
        "full_name": current_user.full_name,
        "role": current_user.role,
    }


@app.get("/dashboard")
def dashboard(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)) -> dict[str, Any]:
    task_counts = {
        row[0]: row[1]
        for row in db.query(QmsTask.status, func.count(QmsTask.id)).group_by(QmsTask.status).all()
    }
    return {
        "documents": db.query(KnowledgeDocument).count(),
        "current_documents": db.query(KnowledgeDocument).filter(KnowledgeDocument.status == "current").count(),
        "tasks": db.query(QmsTask).count(),
        "open_questions": db.query(AgentQuestion).filter(AgentQuestion.status == "Open").count(),
        "awaiting_review": db.query(QmsTask).filter(QmsTask.status == "Awaiting Human Review").count(),
        "task_counts": task_counts,
        "recent_runs": [
            {
                "id": r.id,
                "task_id": r.task_id,
                "run_type": r.run_type,
                "status": r.status,
                "summary": r.summary,
                "started_at": r.started_at,
                "finished_at": r.finished_at,
            }
            for r in db.query(AgentRun).order_by(AgentRun.started_at.desc()).limit(8).all()
        ],
    }


@app.post("/documents")
async def upload_document(
    title: str = Form(...),
    document_number: str = Form(""),
    version: str = Form(""),
    status: str = Form("current"),
    doc_type: str = Form("SOP"),
    department: str = Form(""),
    effective_date: str = Form(""),
    review_date: str = Form(""),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict[str, Any]:
    safe_name = re.sub(r"[^a-zA-Z0-9_.-]", "_", file.filename or "upload")
    stored_name = f"{uuid.uuid4().hex}_{safe_name}"
    path = UPLOAD_DIR / stored_name

    with path.open("wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    text = extract_text_from_upload(path)
    chunks = chunk_text(text)

    doc = KnowledgeDocument(
        title=title,
        document_number=document_number,
        version=version,
        status=status,
        doc_type=doc_type,
        department=department,
        effective_date=parse_date_or_none(effective_date),
        review_date=parse_date_or_none(review_date),
        source_filename=file.filename or safe_name,
        file_path=str(path),
        uploaded_by_id=current_user.id,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    for idx, content in enumerate(chunks):
        db.add(DocumentChunk(document_id=doc.id, chunk_index=idx + 1, content=content))
    db.commit()

    audit(
        db,
        "DOCUMENT_UPLOADED",
        "knowledge_document",
        doc.id,
        current_user,
        {"title": title, "document_number": document_number, "version": version, "chunks": len(chunks)},
    )

    return {"id": doc.id, "title": doc.title, "chunks": len(chunks)}


@app.get("/documents")
def list_documents(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> list[dict[str, Any]]:
    docs = db.query(KnowledgeDocument).order_by(KnowledgeDocument.uploaded_at.desc()).all()
    return [
        {
            "id": d.id,
            "title": d.title,
            "document_number": d.document_number,
            "version": d.version,
            "status": d.status,
            "doc_type": d.doc_type,
            "department": d.department,
            "effective_date": d.effective_date,
            "review_date": d.review_date,
            "source_filename": d.source_filename,
            "uploaded_at": d.uploaded_at,
            "chunks": len(d.chunks),
        }
        for d in docs
    ]


@app.post("/documents/search")
def document_search(
    payload: SearchRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> list[dict[str, Any]]:
    chunks = search_chunks(db, payload.query, payload.include_superseded, limit=12)
    return [
        {
            "document_id": c.document.id,
            "title": c.document.title,
            "document_number": c.document.document_number,
            "version": c.document.version,
            "status": c.document.status,
            "chunk_index": c.chunk_index,
            "content": c.content[:900],
        }
        for c in chunks
    ]


@app.post("/tasks")
def create_task(
    payload: TaskCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict[str, Any]:
    task = QmsTask(
        task_code=f"QA-{utcnow().strftime('%Y%m%d')}-{uuid.uuid4().hex[:6].upper()}",
        title=payload.title,
        task_type=payload.task_type,
        description=payload.description,
        priority=payload.priority,
        due_date=payload.due_date,
        status="Assigned to AI",
        created_by_id=current_user.id,
    )
    db.add(task)
    db.commit()
    db.refresh(task)
    audit(db, "TASK_CREATED", "qms_task", task.id, current_user, {"task_code": task.task_code})
    return serialize_task(task)


@app.get("/tasks")
def list_tasks(
    status: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> list[dict[str, Any]]:
    query = db.query(QmsTask)
    if status:
        query = query.filter(QmsTask.status == status)
    tasks = query.order_by(QmsTask.created_at.desc()).all()
    return [serialize_task(t) for t in tasks]


@app.get("/tasks/{task_id}")
def get_task(
    task_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict[str, Any]:
    task = db.get(QmsTask, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    return serialize_task_detail(db, task)


@app.patch("/tasks/{task_id}")
def update_task(
    task_id: int,
    payload: TaskUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict[str, Any]:
    task = db.get(QmsTask, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")

    for field, value in payload.model_dump(exclude_unset=True).items():
        setattr(task, field, value)
    task.updated_at = utcnow()
    db.commit()
    audit(db, "TASK_UPDATED", "qms_task", task.id, current_user, payload.model_dump(exclude_unset=True))
    return serialize_task(task)


@app.post("/tasks/{task_id}/run-agent")
async def run_task_agent_endpoint(
    task_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict[str, Any]:
    run = await run_agent_for_task(db, task_id, run_type="manual")
    audit(db, "AGENT_RUN_TRIGGERED", "qms_task", task_id, current_user, {"run_id": run.id, "status": run.status})
    task = db.get(QmsTask, task_id)
    return serialize_task_detail(db, task)


@app.post("/automation/run-now")
async def run_automation_now(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict[str, Any]:
    candidates = (
        db.query(QmsTask)
        .filter(QmsTask.status.in_(["New", "Assigned to AI", "In Progress", "Returned for Rework"]))
        .order_by(QmsTask.due_date.asc().nullslast(), QmsTask.created_at.asc())
        .limit(10)
        .all()
    )
    results = []
    for task in candidates:
        task.status = "In Progress"
        task.updated_at = utcnow()
        db.commit()
        run = await run_agent_for_task(db, task.id, run_type="manual_batch")
        results.append({"task_id": task.id, "task_code": task.task_code, "run_status": run.status})
    audit(db, "AUTOMATION_RUN_NOW", "automation", "manual", current_user, {"tasks": len(results)})
    return {"processed": len(results), "results": results}


@app.get("/questions")
def list_questions(
    status: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> list[dict[str, Any]]:
    query = db.query(AgentQuestion).join(QmsTask)
    if status:
        query = query.filter(AgentQuestion.status == status)
    questions = query.order_by(AgentQuestion.created_at.desc()).all()
    return [
        {
            "id": q.id,
            "task_id": q.task_id,
            "task_code": q.task.task_code,
            "task_title": q.task.title,
            "question": q.question,
            "assigned_to": q.assigned_to,
            "status": q.status,
            "answer": q.answer,
            "created_at": q.created_at,
            "answered_at": q.answered_at,
        }
        for q in questions
    ]


@app.post("/questions/{question_id}/answer")
def answer_question(
    question_id: int,
    payload: AnswerQuestionIn,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict[str, Any]:
    q = db.get(AgentQuestion, question_id)
    if not q:
        raise HTTPException(status_code=404, detail="Question not found")
    q.answer = payload.answer
    q.status = "Answered"
    q.answered_at = utcnow()

    remaining = (
        db.query(AgentQuestion)
        .filter(AgentQuestion.task_id == q.task_id, AgentQuestion.status == "Open", AgentQuestion.id != q.id)
        .count()
    )
    if remaining == 0:
        q.task.status = "Assigned to AI"
        q.task.updated_at = utcnow()

    db.commit()
    audit(db, "QUESTION_ANSWERED", "agent_question", q.id, current_user, {"task_id": q.task_id})
    return {"id": q.id, "status": q.status, "task_status": q.task.status}


@app.get("/reviews")
def review_queue(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> list[dict[str, Any]]:
    outputs = (
        db.query(TaskOutput)
        .filter(TaskOutput.status.in_(["Draft Prepared", "Returned for Rework"]))
        .order_by(TaskOutput.created_at.desc())
        .all()
    )
    return [serialize_output(o) for o in outputs]


@app.post("/outputs/{output_id}/review")
def review_output(
    output_id: int,
    payload: ReviewOutputIn,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict[str, Any]:
    output = db.get(TaskOutput, output_id)
    if not output:
        raise HTTPException(status_code=404, detail="Output not found")

    decision = payload.decision.lower().strip()
    if decision not in {"accept", "reject", "rework"}:
        raise HTTPException(status_code=400, detail="Decision must be accept, reject, or rework")

    output.reviewer_comment = payload.comment
    output.reviewed_at = utcnow()

    if decision == "accept":
        output.status = "Human Accepted"
        output.task.status = "Human Approved"
    elif decision == "reject":
        output.status = "Human Rejected"
        output.task.status = "Closed - Rejected"
    else:
        output.status = "Returned for Rework"
        output.task.status = "Returned for Rework"

    output.task.updated_at = utcnow()
    db.commit()
    audit(
        db,
        "OUTPUT_REVIEWED",
        "task_output",
        output.id,
        current_user,
        {"decision": decision, "task_id": output.task_id, "comment": payload.comment},
    )
    return serialize_output(output)


@app.get("/audit")
def list_audit_events(
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> list[dict[str, Any]]:
    events = db.query(AuditEvent).order_by(AuditEvent.created_at.desc()).limit(min(limit, 500)).all()
    return [
        {
            "id": e.id,
            "event_type": e.event_type,
            "entity_type": e.entity_type,
            "entity_id": e.entity_id,
            "user_id": e.user_id,
            "details": json.loads(e.details_json or "{}"),
            "created_at": e.created_at,
        }
        for e in events
    ]


def serialize_task(task: QmsTask) -> dict[str, Any]:
    return {
        "id": task.id,
        "task_code": task.task_code,
        "title": task.title,
        "task_type": task.task_type,
        "description": task.description,
        "priority": task.priority,
        "status": task.status,
        "due_date": task.due_date,
        "created_at": task.created_at,
        "updated_at": task.updated_at,
        "last_agent_run_at": task.last_agent_run_at,
    }


def serialize_output(output: TaskOutput) -> dict[str, Any]:
    return {
        "id": output.id,
        "task_id": output.task_id,
        "task_code": output.task.task_code,
        "task_title": output.task.title,
        "title": output.title,
        "content": output.content,
        "sources": json.loads(output.sources_json or "[]"),
        "assumptions": json.loads(output.assumptions_json or "[]"),
        "status": output.status,
        "reviewer_comment": output.reviewer_comment,
        "created_at": output.created_at,
        "reviewed_at": output.reviewed_at,
    }


def serialize_task_detail(db: Session, task: QmsTask) -> dict[str, Any]:
    base = serialize_task(task)
    base["questions"] = [
        {
            "id": q.id,
            "question": q.question,
            "assigned_to": q.assigned_to,
            "status": q.status,
            "answer": q.answer,
            "created_at": q.created_at,
            "answered_at": q.answered_at,
        }
        for q in task.questions
    ]
    base["outputs"] = [serialize_output(o) for o in task.outputs]
    base["agent_runs"] = [
        {
            "id": r.id,
            "run_type": r.run_type,
            "status": r.status,
            "summary": r.summary,
            "sources": json.loads(r.sources_json or "[]"),
            "started_at": r.started_at,
            "finished_at": r.finished_at,
        }
        for r in db.query(AgentRun).filter(AgentRun.task_id == task.id).order_by(AgentRun.started_at.desc()).all()
    ]
    return base
